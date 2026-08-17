from __future__ import annotations

import base64
import html
import io
import json
import re
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

import arabic_reshaper
from bidi.algorithm import get_display

LABELS = {
    "ar": {"summary": "الملخص المهني", "experience": "الخبرة المهنية", "education": "التعليم", "skills": "المهارات", "languages": "اللغات", "extras": "معلومات إضافية"},
    "fr": {"summary": "Profil professionnel", "experience": "Expérience professionnelle", "education": "Formation", "skills": "Compétences", "languages": "Langues", "extras": "Informations complémentaires"},
    "en": {"summary": "Professional Summary", "experience": "Work Experience", "education": "Education", "skills": "Skills", "languages": "Languages", "extras": "Additional Information"},
}
AR_RE = re.compile(r"[؀-ۿ]")

def _items(data: dict, key: str) -> list[dict]:
    return [item for item in data.get(key, []) if isinstance(item, dict) and any(str(v).strip() for v in item.values())]

def _photo_bytes(data: dict) -> bytes | None:
    value = str(data.get("photoData", ""))
    match = re.fullmatch(r"data:image/(?:png|jpe?g|webp);base64,([A-Za-z0-9+/=\s]+)", value)
    if not match: return None
    try: raw = base64.b64decode(match.group(1), validate=True)
    except Exception: return None
    return raw if raw and len(raw) > 0 else None

def _lines(data: dict) -> Iterable[tuple[str, list[str]]]:
    labels = LABELS.get(data.get("uiLanguage", "ar"), LABELS["ar"])
    yield labels["summary"], [str(data.get("summary", ""))]
    exp = []
    for x in _items(data, "experience"):
        exp.extend([f'{x.get("role", "")} — {x.get("company", "")}', f'{x.get("start", "")} – {x.get("end", "")}', str(x.get("details", ""))])
    yield labels["experience"], exp
    edu = []
    for x in _items(data, "education"):
        edu.extend([f'{x.get("degree", "")} — {x.get("institution", "")}', f'{x.get("start", "")} – {x.get("end", "")}', str(x.get("details", ""))])
    yield labels["education"], edu
    yield labels["skills"], [str(data.get("skills", ""))]
    yield labels["languages"], [str(data.get("languages", ""))]
    yield labels["extras"], [str(data.get("extras", ""))]

def export_txt(data: dict) -> bytes:
    contact = " | ".join(filter(None, [data.get("email"), data.get("phone"), data.get("location"), data.get("linkedin")]))
    output = [str(data.get("fullName", "")), str(data.get("headline", "")), contact]
    for title, lines in _lines(data):
        clean = [l.strip() for l in lines if l and l.strip()]
        if clean: output.extend(["", title, *clean])
    return "\n".join(output).encode("utf-8-sig")

def export_json(data: dict) -> bytes:
    return json.dumps(data, ensure_ascii=False, indent=2).encode("utf-8-sig")

def export_html(data: dict) -> bytes:
    lang = data.get("uiLanguage", "ar")
    direction = "rtl" if lang == "ar" else "ltr"
    photo = str(data.get("photoData", ""))
    photo_html = f'<img class="photo" src="{photo}" alt="photo" />' if _photo_bytes(data) else ""
    sections = []
    for title, lines in _lines(data):
        clean = [html.escape(l).replace("\n", " ") for l in lines if l and l.strip()]
        if clean:
            sections.append(f'<section><h2>{html.escape(title)}</h2>' + "".join(f"<p>{l}</p>" for l in clean) + "</section>")
    contact = " · ".join(html.escape(str(v)) for v in [data.get("email"), data.get("phone"), data.get("location"), data.get("linkedin")] if v)
    page = f'''<!doctype html><html lang="{lang}" dir="{direction}"><head><meta charset="utf-8" /><title>{html.escape(str(data.get("fullName", "")))}</title>
<style>body{{font-family:Arial,Tahoma,sans-serif;margin:32px;color:#0f172a}}.photo{{width:96px;height:96px;border-radius:50%;object-fit:cover}}h1{{margin:8px 0 2px}}.contact{{color:#475569;margin-bottom:16px}}h2{{color:#1e745a;border-bottom:2px solid #e2e8f0;padding-bottom:4px}}</style>
</head><body>{photo_html}<h1>{html.escape(str(data.get("fullName", "")))}</h1><p>{html.escape(str(data.get("headline", "")))}</p><p class="contact">{contact}</p>{"".join(sections)}</body></html>'''
    return page.encode("utf-8-sig")

def export_docx(data: dict) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(10.5)
    align = WD_ALIGN_PARAGRAPH.RIGHT if data.get("uiLanguage", "ar") == "ar" else WD_ALIGN_PARAGRAPH.LEFT
    photo = _photo_bytes(data)
    if photo:
        p = doc.add_paragraph(); p.alignment = align; p.add_run().add_picture(io.BytesIO(photo), width=Inches(1.05))
    title = doc.add_heading(str(data.get("fullName", "")), 0); title.alignment = align
    for text in [data.get("headline"), " | ".join(filter(None, [data.get("email"), data.get("phone"), data.get("location"), data.get("linkedin")]))]:
        if text:
            p = doc.add_paragraph(str(text)); p.alignment = align
    for heading, lines in _lines(data):
        clean = [l.strip() for l in lines if l and l.strip()]
        if not clean: continue
        h = doc.add_heading(heading, 1); h.alignment = align
        for line in clean:
            p = doc.add_paragraph(line); p.alignment = align
    buffer = io.BytesIO(); doc.save(buffer); return buffer.getvalue()

def export_pdf(data: dict) -> bytes:
    buffer = io.BytesIO()
    font = "Helvetica"
    for path in ["/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", "C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/tahoma.ttf"]:
        try: pdfmetrics.registerFont(TTFont("CvUnicode", path)); font = "CvUnicode"; break
        except Exception: continue
    is_arabic = data.get("uiLanguage", "ar") == "ar"
    styles = getSampleStyleSheet()
    body = ParagraphStyle("CvBody", parent=styles["BodyText"], fontName=font, fontSize=10, leading=15, alignment=TA_RIGHT if is_arabic else TA_LEFT)
    heading = ParagraphStyle("CvHeading", parent=body, fontSize=14, leading=20, textColor="#1e745a", spaceBefore=8)
    title = ParagraphStyle("CvTitle", parent=body, fontSize=22, leading=28)
    def pdf_text(value: object) -> str:
        raw = str(value or "")
        if is_arabic and AR_RE.search(raw): raw = get_display(arabic_reshaper.reshape(raw), base_dir="R")
        return html.escape(raw).replace("\n", " ")
    story = []
    photo = _photo_bytes(data)
    if photo: story.extend([PdfImage(io.BytesIO(photo), width=26 * mm, height=26 * mm), Spacer(1, 5)])
    story.append(Paragraph(pdf_text(data.get("fullName", "")), title))
    contact_rtl = " | ".join(filter(None, [data.get("headline"), data.get("location")]))
    contact_ltr = " | ".join(filter(None, [data.get("email"), data.get("phone"), data.get("linkedin")]))
    if contact_rtl: story.append(Paragraph(pdf_text(contact_rtl), body))
    if contact_ltr: story.append(Paragraph(html.escape(contact_ltr).replace("\n", " "), body))
    if contact_rtl or contact_ltr: story.append(Spacer(1, 5))
    for section, lines in _lines(data):
        clean = [l.strip() for l in lines if l and l.strip()]
        if clean:
            story.append(Paragraph(pdf_text(section), heading))
            story.extend(Paragraph(pdf_text(line), body) for line in clean)
    doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    doc.build(story)
    return buffer.getvalue()